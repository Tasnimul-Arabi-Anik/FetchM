from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

import pandas as pd
from scipy import stats

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional at runtime until dependency is installed
    Document = None


MISSING_LABELS = {"absent", "unknown"}


def _safe_top_counts(series: pd.Series, *, exclude: Sequence[str] = (), limit: int = 10) -> List[tuple[str, int]]:
    filtered = series.dropna()
    if exclude:
        filtered = filtered[~filtered.isin(list(exclude))]
    counts = filtered.value_counts().head(limit)
    return [(str(index), int(value)) for index, value in counts.items()]


def _format_top_counts(items: Iterable[tuple[str, int]]) -> List[str]:
    return [f"{label}: {count}" for label, count in items]


def format_duration(runtime_seconds: float) -> str:
    if runtime_seconds < 60:
        return f"{runtime_seconds:.1f} seconds"
    if runtime_seconds < 3600:
        return f"{runtime_seconds / 60:.2f} minutes"
    return f"{runtime_seconds / 3600:.2f} hours"


def _format_count_pct(count: int, total: int) -> str:
    pct = 0.0 if total == 0 else (count / total) * 100
    return f"{pct:.2f}% (n = {count:,})"


def _format_numeric(value: float, *, decimals: int = 2) -> str:
    return f"{value:,.{decimals}f}"


def _format_pvalue(value: float) -> str:
    if value < 0.001:
        return "< 0.001"
    return f"{value:.3f}"


def _clean_counts(series: pd.Series, *, limit: int = 10) -> List[tuple[str, int]]:
    return _safe_top_counts(series, exclude=tuple(MISSING_LABELS), limit=limit)


def _summarize_numeric(series: pd.Series) -> Dict[str, float] | None:
    values = pd.to_numeric(series, errors="coerce").dropna()
    if values.empty:
        return None
    return {
        "min": float(values.min()),
        "median": float(values.median()),
        "mean": float(values.mean()),
        "max": float(values.max()),
    }


def _build_field_insight(field: str, present: int, unknown: int, absent: int, total: int) -> str:
    present_pct = 0.0 if total == 0 else (present / total) * 100
    unknown_pct = 0.0 if total == 0 else (unknown / total) * 100
    absent_pct = 0.0 if total == 0 else (absent / total) * 100
    return (
        f"{field} was informative for {present} records ({present_pct:.2f}%), "
        f"explicitly unknown in {unknown} records ({unknown_pct:.2f}%), "
        f"and absent after retrieval in {absent} records ({absent_pct:.2f}%)."
    )


def _build_dataset_insights(
    df_clean: pd.DataFrame,
    completeness_rows: List[Dict[str, object]],
    runtime_seconds: float,
    total_input_rows: int,
    processed_rows: int,
    unique_assemblies: int,
) -> List[str]:
    insights = [
        (
            f"FetchM processed {processed_rows} rows from {total_input_rows} NCBI input rows and produced "
            f"{unique_assemblies} unique assemblies in {format_duration(runtime_seconds)}."
        )
    ]

    field_map = {row["field"]: row for row in completeness_rows}
    for key in ["Collection Date", "Geographic Location", "Host"]:
        row = field_map.get(key)
        if row:
            insights.append(
                _build_field_insight(
                    key,
                    int(row["present"]),
                    int(row["unknown"]),
                    int(row["absent"]),
                    len(df_clean),
                )
            )

    top_locations = _clean_counts(df_clean["Geographic Location"], limit=3)
    if top_locations:
        formatted = ", ".join(f"{name} ({count})" for name, count in top_locations)
        insights.append(f"The most represented geographic locations in the cleaned dataset were {formatted}.")

    top_hosts = _clean_counts(df_clean["Host"], limit=3)
    if top_hosts:
        formatted = ", ".join(f"{name} ({count})" for name, count in top_hosts)
        insights.append(f"The most represented host annotations were {formatted}.")

    valid_years = pd.to_numeric(
        df_clean.loc[~df_clean["Collection Date"].isin(list(MISSING_LABELS)), "Collection Date"],
        errors="coerce",
    ).dropna()
    if not valid_years.empty:
        insights.append(
            f"Collection years spanned {int(valid_years.min())} to {int(valid_years.max())}, "
            f"supporting temporal analyses across the retained assemblies."
        )

    return insights


def _compute_correlation_summary(x: pd.Series, y: pd.Series, label: str) -> Dict[str, object] | None:
    x_numeric = pd.to_numeric(x, errors="coerce").round().astype("Int64")
    y_numeric = pd.to_numeric(y, errors="coerce")
    plot_data = pd.DataFrame({"x": x_numeric, "y": y_numeric}).dropna()
    if len(plot_data) < 2:
        return None

    q1 = plot_data["x"].quantile(0.25)
    q3 = plot_data["x"].quantile(0.75)
    iqr = q3 - q1
    lower_bound = q1 - 1.5 * iqr
    upper_bound = q3 + 1.5 * iqr
    filtered = plot_data[(plot_data["x"] >= lower_bound) & (plot_data["x"] <= upper_bound)]
    if len(filtered) < 2:
        return None
    if filtered["x"].nunique() < 2 or filtered["y"].nunique() < 2:
        return None

    pearson_r, pearson_p = stats.pearsonr(filtered["x"], filtered["y"])
    spearman_r, spearman_p = stats.spearmanr(filtered["x"], filtered["y"])
    return {
        "label": label,
        "n": int(len(filtered)),
        "pearson_r": float(pearson_r),
        "pearson_p": float(pearson_p),
        "spearman_rho": float(spearman_r),
        "spearman_p": float(spearman_p),
    }


def _build_trend_summaries(df_clean: pd.DataFrame, df_annotation: pd.DataFrame) -> List[Dict[str, object]]:
    summaries: List[Dict[str, object]] = []
    summaries_item = [
        (
            "Assembly sequence length",
            df_clean["Collection Date"],
            df_clean["Assembly Stats Total Sequence Length"],
        ),
        (
            "Total annotated genes",
            df_annotation["Collection Date"],
            df_annotation["Annotation Count Gene Total"],
        ),
        (
            "Protein-coding genes",
            df_annotation["Collection Date"],
            df_annotation["Annotation Count Gene Protein-coding"],
        ),
        (
            "Pseudogenes",
            df_annotation["Collection Date"],
            df_annotation["Annotation Count Gene Pseudogene"],
        ),
    ]
    for label, x, y in summaries_item:
        summary = _compute_correlation_summary(x, y, label)
        if summary:
            summaries.append(summary)
    return summaries


def _build_filter_narrative(filter_summary: Dict[str, object]) -> str:
    retained_rows = int(filter_summary["retained_rows"])
    total_rows = int(filter_summary["total_input_rows"])
    removed_rows = int(filter_summary["total_removed_rows"])

    clauses = [
        (
            f"Using FetchM, {retained_rows:,} NCBI input rows were retained after processing from "
            f"{total_rows:,} total input rows."
        )
    ]
    if filter_summary["checkm_enabled"]:
        clauses.append(
            f"CheckM completeness filtering was applied at > {filter_summary['checkm_threshold']}, "
            f"removing {int(filter_summary['checkm_removed_rows']):,} rows."
        )
    else:
        clauses.append("CheckM completeness filtering was not applied.")

    if filter_summary["ani_enabled"]:
        ani_values = ", ".join(str(value) for value in filter_summary["ani_values"] if value != "all")
        clauses.append(
            f"ANI filtering retained rows with status {ani_values}, removing "
            f"{int(filter_summary['ani_removed_rows']):,} rows."
        )
    else:
        clauses.append("ANI filtering was disabled.")

    clauses.append(f"The total number of rows removed by filtering was {removed_rows:,}.")
    return " ".join(clauses)


def _build_numeric_paragraphs(context: Dict[str, object]) -> List[str]:
    paragraphs: List[str] = []
    numeric_summary = context["numeric_summary"]

    assembly_summary = numeric_summary.get("Assembly Stats Total Sequence Length")
    if assembly_summary:
        paragraphs.append(
            "The total sequence length ranged from "
            f"{_format_numeric(assembly_summary['min'])} bp to {_format_numeric(assembly_summary['max'])} bp, "
            f"with a mean of {_format_numeric(assembly_summary['mean'])} bp and a median of "
            f"{_format_numeric(assembly_summary['median'])} bp."
        )

    gene_summary = numeric_summary.get("Annotation Count Gene Total")
    protein_summary = numeric_summary.get("Annotation Count Gene Protein-coding")
    pseudogene_summary = numeric_summary.get("Annotation Count Gene Pseudogene")
    pgap_rows = int(context["pgap_annotation_rows"])
    if gene_summary or protein_summary or pseudogene_summary:
        lead = (
            f"Gene annotation metrics were summarized for {pgap_rows:,} rows annotated with the "
            "NCBI Prokaryotic Genome Annotation Pipeline (PGAP)."
        )
        details: List[str] = []
        if gene_summary:
            details.append(
                "The total number of annotated genes ranged from "
                f"{_format_numeric(gene_summary['min'])} to {_format_numeric(gene_summary['max'])}, with a mean of "
                f"{_format_numeric(gene_summary['mean'])} and a median of {_format_numeric(gene_summary['median'])}."
            )
        if protein_summary:
            details.append(
                "Protein-coding gene counts ranged from "
                f"{_format_numeric(protein_summary['min'])} to {_format_numeric(protein_summary['max'])}, with a mean of "
                f"{_format_numeric(protein_summary['mean'])} and a median of {_format_numeric(protein_summary['median'])}."
            )
        if pseudogene_summary:
            details.append(
                "Pseudogene counts ranged from "
                f"{_format_numeric(pseudogene_summary['min'])} to {_format_numeric(pseudogene_summary['max'])}, with a mean of "
                f"{_format_numeric(pseudogene_summary['mean'])} and a median of {_format_numeric(pseudogene_summary['median'])}."
            )
        paragraphs.append(" ".join([lead] + details))

    return paragraphs


def _build_temporal_paragraph(context: Dict[str, object]) -> str:
    df_clean = context["df_clean"]
    total = len(df_clean)
    present = int((~df_clean["Collection Date"].isin(list(MISSING_LABELS))).sum())
    unknown = int((df_clean["Collection Date"] == "unknown").sum())
    absent = int((df_clean["Collection Date"] == "absent").sum())
    missing_total = unknown + absent
    valid_years = pd.to_numeric(
        df_clean.loc[~df_clean["Collection Date"].isin(list(MISSING_LABELS)), "Collection Date"],
        errors="coerce",
    ).dropna()

    paragraph = (
        f"Temporal metadata was available for {_format_count_pct(present, total)} of retained records, while "
        f"{_format_count_pct(missing_total, total)} lacked a usable collection year. "
        f"Of the missing values, {_format_count_pct(unknown, total)} were explicitly marked as unknown in the source "
        f"metadata and {_format_count_pct(absent, total)} were absent after retrieval."
    )
    if not valid_years.empty:
        peak_year = int(valid_years.value_counts().idxmax())
        peak_count = int(valid_years.value_counts().max())
        paragraph += (
            f" Collection years ranged from {int(valid_years.min())} to {int(valid_years.max())}, "
            f"with the highest count in {peak_year} at {_format_count_pct(peak_count, total)}."
        )
    return paragraph


def _top_distribution_sentence(series: pd.Series, total: int, label: str, limit: int = 5) -> str:
    items = _clean_counts(series, limit=limit)
    if not items:
        return f"No informative {label.lower()} values were available."
    formatted = ", ".join(f"{name} ({_format_count_pct(count, total)})" for name, count in items)
    return f"The highest contributions for {label.lower()} were {formatted}."


def _build_geography_paragraphs(context: Dict[str, object]) -> List[str]:
    df_clean = context["df_clean"]
    total = len(df_clean)
    informative_geo = df_clean.loc[~df_clean["Geographic Location"].isin(list(MISSING_LABELS)), "Geographic Location"]
    geo_unique = int(informative_geo.nunique())
    geo_absent = int((df_clean["Geographic Location"] == "absent").sum())
    geo_unknown = int((df_clean["Geographic Location"] == "unknown").sum())
    geo_missing = geo_absent + geo_unknown

    paragraphs = [
        (
            f"Geographic metadata analysis identified {geo_unique:,} unique geographic locations in the retained dataset. "
            + _top_distribution_sentence(df_clean["Geographic Location"], total, "geographic location")
        ),
        _top_distribution_sentence(df_clean["Subcontinent"], total, "subcontinent"),
        (
            _top_distribution_sentence(df_clean["Continent"], total, "continent")
            + f" Geographic metadata remained incomplete for {_format_count_pct(geo_missing, total)} of retained records, "
            f"including {_format_count_pct(geo_unknown, total)} explicitly unknown values and "
            f"{_format_count_pct(geo_absent, total)} absent-after-retrieval values."
        ),
    ]
    return paragraphs


def _build_host_paragraph(context: Dict[str, object]) -> str:
    df_clean = context["df_clean"]
    total = len(df_clean)
    host_absent = int((df_clean["Host"] == "absent").sum())
    host_unknown = int((df_clean["Host"] == "unknown").sum())
    host_present = total - host_absent - host_unknown
    paragraph = (
        f"Host metadata was informative for {_format_count_pct(host_present, total)} of retained records. "
        f"A further {_format_count_pct(host_unknown, total)} were explicitly unknown and "
        f"{_format_count_pct(host_absent, total)} were absent after retrieval."
    )
    top_hosts = _clean_counts(df_clean["Host"], limit=5)
    if top_hosts:
        formatted = ", ".join(f"{name} ({_format_count_pct(count, total)})" for name, count in top_hosts)
        paragraph += f" The most frequent host annotations were {formatted}."
    return paragraph


def _build_trend_paragraphs(context: Dict[str, object]) -> List[str]:
    paragraphs: List[str] = []
    for trend in context["trend_summaries"]:
        paragraphs.append(
            f"{trend['label']} over collection year was evaluated on {trend['n']:,} usable records after the same "
            f"outlier filtering used for the scatter plots. Pearson's r was {trend['pearson_r']:.3f} "
            f"(p {_format_pvalue(trend['pearson_p'])}) and Spearman's rho was {trend['spearman_rho']:.3f} "
            f"(p {_format_pvalue(trend['spearman_p'])})."
        )
    return paragraphs


def _build_validation_paragraph(context: Dict[str, object]) -> str:
    total_rows = int(context["processed_rows"])
    ok_like = int(context["fetch_status_counts"].get("ok", 0)) + int(context["fetch_status_counts"].get("cached", 0))
    failure_total = int(context["fetch_failure_total"])
    unique_biosamples = int(context["unique_biosamples"])
    failure_counts = context["fetch_failure_counts"]

    paragraph = (
        f"Validation and data-quality auditing were based on {total_rows:,} processed rows linked to "
        f"{unique_biosamples:,} unique BioSample accessions. Metadata retrieval completed without recorded fetch "
        f"failure for {ok_like:,} rows, while {failure_total:,} rows were flagged in the fetch-failure audit."
    )
    if failure_counts:
        reasons = ", ".join(f"{reason} ({count:,})" for reason, count in failure_counts.items())
        paragraph += f" Recorded failure reasons were {reasons}."
    paragraph += (
        " FetchM now separates source-explicit missing values (`unknown`) from values that remained unavailable "
        "after retrieval (`absent`), which helps distinguish source curation limits from retrieval gaps."
    )
    return paragraph


def _build_runtime_paragraph(context: Dict[str, object]) -> str:
    sentence = (
        f"The {context['mode']} workflow completed in {context['runtime_human']} for "
        f"{int(context['total_input_rows']):,} NCBI input rows, with {int(context['processed_rows']):,} rows "
        "retained after filtering."
    )
    if context["mode"] == "run":
        sentence += f" Sequence downloading produced {int(context['sequence_download_count']):,} FASTA files."
    sentence += " Runtime will vary with dataset size, network conditions, and NCBI responsiveness."
    return sentence


def _build_narrative_sections(context: Dict[str, object]) -> List[Dict[str, object]]:
    sections = [
        {
            "heading": "Narrative Results",
            "paragraphs": [
                _build_filter_narrative(context["filter_summary"]),
                *_build_numeric_paragraphs(context),
                _build_temporal_paragraph(context),
                *_build_geography_paragraphs(context),
                _build_host_paragraph(context),
            ]
            + _build_trend_paragraphs(context)
            + [
                _build_runtime_paragraph(context),
            ],
        },
        {
            "heading": "Validation and Data Quality",
            "paragraphs": [
                _build_validation_paragraph(context),
            ],
        },
    ]
    return sections


def build_report_context(
    *,
    organism_name: str,
    input_file: str,
    output_root: str,
    mode: str,
    total_input_rows: int,
    processed_rows: int,
    unique_assemblies: int,
    runtime_seconds: float,
    filters: Dict[str, str],
    filter_summary: Dict[str, object],
    df_clean: pd.DataFrame,
    df_annotation: pd.DataFrame,
    fetch_status_df: pd.DataFrame,
) -> Dict[str, object]:
    completeness_rows = []
    for column in ["Isolation Source", "Collection Date", "Geographic Location", "Host", "Continent", "Subcontinent"]:
        absent_count = int((df_clean[column] == "absent").sum())
        unknown_count = int((df_clean[column] == "unknown").sum())
        present_count = len(df_clean) - absent_count - unknown_count
        completeness_rows.append(
            {
                "field": column,
                "present": present_count,
                "unknown": unknown_count,
                "absent": absent_count,
            }
        )

    numeric_summary = {}
    for column in [
        "Assembly Stats Total Sequence Length",
        "Annotation Count Gene Total",
        "Annotation Count Gene Protein-coding",
        "Annotation Count Gene Pseudogene",
    ]:
        source_df = df_annotation if column.startswith("Annotation") else df_clean
        summary = _summarize_numeric(source_df[column])
        if summary:
            numeric_summary[column] = summary

    figures_dir = Path(output_root) / "figures"
    metadata_dir = Path(output_root) / "metadata_output"
    sequence_dir = Path(output_root) / "sequence"

    sequence_fasta_files = sorted(path.name for path in sequence_dir.glob("*.fna"))
    dataset_insights = _build_dataset_insights(
        df_clean=df_clean,
        completeness_rows=completeness_rows,
        runtime_seconds=runtime_seconds,
        total_input_rows=total_input_rows,
        processed_rows=processed_rows,
        unique_assemblies=unique_assemblies,
    )
    if "Metadata Fetch Status" in fetch_status_df.columns and "Metadata Fetch Reason" in fetch_status_df.columns:
        failure_only = fetch_status_df[fetch_status_df["Metadata Fetch Status"] != "ok"].copy()
        failure_only = failure_only[failure_only["Metadata Fetch Status"] != "cached"]
        fetch_failure_total = len(failure_only)
        fetch_failure_counts = failure_only["Metadata Fetch Reason"].fillna("missing_reason").value_counts().to_dict()
        fetch_status_counts = fetch_status_df["Metadata Fetch Status"].fillna("missing_status").value_counts().to_dict()
    else:
        fetch_failure_total = 0
        fetch_failure_counts = {}
        fetch_status_counts = {}

    context = {
        "title": f"FetchM Comprehensive Report: {organism_name}",
        "organism_name": organism_name,
        "mode": mode,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "input_file": input_file,
        "output_root": output_root,
        "runtime_human": format_duration(runtime_seconds),
        "total_input_rows": total_input_rows,
        "processed_rows": processed_rows,
        "unique_assemblies": unique_assemblies,
        "filters": filters,
        "filter_summary": filter_summary,
        "completeness_rows": completeness_rows,
        "top_years": _format_top_counts(_clean_counts(df_clean["Collection Date"])),
        "top_locations": _format_top_counts(_clean_counts(df_clean["Geographic Location"])),
        "top_hosts": _format_top_counts(_clean_counts(df_clean["Host"])),
        "numeric_summary": numeric_summary,
        "dataset_insights": dataset_insights,
        "fetch_failure_total": fetch_failure_total,
        "fetch_failure_counts": fetch_failure_counts,
        "fetch_status_counts": fetch_status_counts,
        "unique_biosamples": int(fetch_status_df["Assembly BioSample Accession"].dropna().astype(str).nunique())
        if "Assembly BioSample Accession" in fetch_status_df.columns
        else 0,
        "sequence_download_count": len(sequence_fasta_files),
        "figure_files": sorted(path.name for path in figures_dir.glob("*") if path.is_file()),
        "metadata_files": sorted(path.name for path in metadata_dir.glob("*") if path.is_file()),
        "sequence_files": sorted(path.name for path in sequence_dir.glob("*") if path.is_file()),
        "pgap_annotation_rows": int(len(df_annotation)),
        "trend_summaries": _build_trend_summaries(df_clean, df_annotation),
        "df_clean": df_clean,
    }
    context["narrative_sections"] = _build_narrative_sections(context)
    return context


def render_markdown_report(context: Dict[str, object], output_path: str) -> None:
    lines = [
        f"# {context['title']}",
        "",
        "## Run Summary",
        f"- Target organism: {context['organism_name']}",
        f"- Workflow mode: {context['mode']}",
        f"- Report generated: {context['created_at']}",
        f"- Input file: `{context['input_file']}`",
        f"- Output directory: `{context['output_root']}`",
        f"- Runtime: {context['runtime_human']}",
        f"- Total NCBI input rows: {context['total_input_rows']}",
        f"- Rows processed after filtering: {context['processed_rows']}",
        f"- Unique assemblies in clean output: {context['unique_assemblies']}",
        "",
        "## Active Filters and Runtime Settings",
    ]

    for key, value in context["filters"].items():
        lines.append(f"- {key}: {value}")

    for section in context["narrative_sections"]:
        lines.extend(["", f"## {section['heading']}"])
        for paragraph in section["paragraphs"]:
            lines.extend([paragraph, ""])
        if lines[-1] == "":
            lines.pop()

    lines.extend(["", "## Metadata Completeness"])
    for row in context["completeness_rows"]:
        lines.append(
            f"- {row['field']}: present={row['present']}, unknown={row['unknown']}, absent={row['absent']}"
        )

    lines.extend(["", "## Key Observations"])
    for insight in context["dataset_insights"]:
        lines.append(f"- {insight}")

    lines.extend(["", "## Distribution Highlights"])
    for heading, values in [
        ("Top collection years", context["top_years"]),
        ("Top geographic locations", context["top_locations"]),
        ("Top host values", context["top_hosts"]),
    ]:
        lines.append(f"### {heading}")
        if values:
            lines.extend([f"- {value}" for value in values])
        else:
            lines.append("- No informative values available.")

    lines.extend(
        [
            "",
            "## Interpretation",
            (
                "The report distinguishes two classes of incomplete metadata. "
                "`unknown` indicates that the source record explicitly used a missing or unknown-style value, "
                "whereas `absent` indicates that FetchM could not retrieve or locate a usable value for that field."
            ),
            (
                "This distinction is intended to separate limitations in source curation from limitations caused by "
                "missing linked metadata or retrieval failure, which is important for downstream interpretation."
            ),
        ]
    )

    lines.extend(["", "## Fetch Failure Summary"])
    lines.append(f"- Metadata fetch records not marked `ok`: {context['fetch_failure_total']}")
    if context["fetch_failure_counts"]:
        for reason, count in context["fetch_failure_counts"].items():
            lines.append(f"- {reason}: {count}")
    else:
        lines.append("- No fetch failures recorded.")

    lines.extend(
        [
            "",
            "## Overall Conclusion",
            (
                "FetchM generated a cleaned, analysis-ready dataset and a reproducible metadata summary for the target "
                "organism. The remaining gaps are now explicitly separated into source-side unknown values and fields "
                "that were absent after retrieval, helping users interpret both metadata quality and retrieval limits."
            ),
        ]
    )

    if context["mode"] == "run":
        lines.extend(
            [
                "",
                "## Sequence Download Summary",
                f"- Downloaded FASTA files detected: {context['sequence_download_count']}",
            ]
        )

    lines.extend(["", "## Numeric Summaries"])
    for column, summary in context["numeric_summary"].items():
        lines.append(
            f"- {column}: min={summary['min']:.2f}, median={summary['median']:.2f}, "
            f"mean={summary['mean']:.2f}, max={summary['max']:.2f}"
        )

    lines.extend(["", "## Generated Output Files"])
    for section, values in [
        ("Metadata files", context["metadata_files"]),
        ("Figure files", context["figure_files"]),
        ("Sequence files", context["sequence_files"]),
    ]:
        lines.append(f"### {section}")
        if values:
            lines.extend([f"- {value}" for value in values])
        else:
            lines.append("- None")

    Path(output_path).write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_docx_report(context: Dict[str, object], output_path: str) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to generate DOCX reports.")

    document = Document()
    document.add_heading(context["title"], level=0)
    document.add_paragraph(
        f"FetchM generated this report for the target organism {context['organism_name']} "
        f"using the {context['mode']} workflow."
    )

    document.add_heading("Run Summary", level=1)
    for label, value in [
        ("Report generated", context["created_at"]),
        ("Input file", context["input_file"]),
        ("Output directory", context["output_root"]),
        ("Runtime", context["runtime_human"]),
        ("Total NCBI input rows", context["total_input_rows"]),
        ("Rows processed after filtering", context["processed_rows"]),
        ("Unique assemblies in clean output", context["unique_assemblies"]),
    ]:
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Active Filters and Runtime Settings", level=1)
    for key, value in context["filters"].items():
        document.add_paragraph(f"{key}: {value}", style="List Bullet")

    for section in context["narrative_sections"]:
        document.add_heading(section["heading"], level=1)
        for paragraph in section["paragraphs"]:
            document.add_paragraph(paragraph)

    document.add_heading("Metadata Completeness", level=1)
    table = document.add_table(rows=1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Present"
    header_cells[2].text = "Unknown"
    header_cells[3].text = "Absent"
    for row in context["completeness_rows"]:
        cells = table.add_row().cells
        cells[0].text = row["field"]
        cells[1].text = str(row["present"])
        cells[2].text = str(row["unknown"])
        cells[3].text = str(row["absent"])

    document.add_heading("Key Observations", level=1)
    for insight in context["dataset_insights"]:
        document.add_paragraph(insight, style="List Bullet")

    document.add_heading("Distribution Highlights", level=1)
    for heading, values in [
        ("Top collection years", context["top_years"]),
        ("Top geographic locations", context["top_locations"]),
        ("Top host values", context["top_hosts"]),
    ]:
        document.add_paragraph(heading, style="List Bullet")
        if values:
            for value in values:
                document.add_paragraph(value, style="List Bullet 2")
        else:
            document.add_paragraph("No informative values available.", style="List Bullet 2")

    document.add_heading("Interpretation", level=1)
    document.add_paragraph(
        "This report distinguishes two forms of incomplete metadata. "
        "'unknown' denotes records where the source metadata explicitly used a missing or unknown-style value. "
        "'absent' denotes records where FetchM could not retrieve or identify a usable value for that field."
    )
    document.add_paragraph(
        "This separation is intended to help users distinguish source-side curation limits from retrieval-side gaps "
        "when interpreting metadata completeness and downstream analyses."
    )

    document.add_heading("Fetch Failure Summary", level=1)
    document.add_paragraph(
        f"Metadata fetch records not marked 'ok': {context['fetch_failure_total']}"
    )
    if context["fetch_failure_counts"]:
        for reason, count in context["fetch_failure_counts"].items():
            document.add_paragraph(f"{reason}: {count}", style="List Bullet")
    else:
        document.add_paragraph("No fetch failures recorded.", style="List Bullet")

    document.add_heading("Overall Conclusion", level=1)
    document.add_paragraph(
        "FetchM generated a cleaned, analysis-ready dataset and a reproducible metadata summary for the target "
        "organism. The remaining gaps are explicitly separated into source-side unknown values and fields that were "
        "absent after retrieval, helping users interpret both metadata quality and retrieval limits."
    )

    if context["mode"] == "run":
        document.add_heading("Sequence Download Summary", level=1)
        document.add_paragraph(
            f"Downloaded FASTA files detected: {context['sequence_download_count']}"
        )

    document.add_heading("Numeric Summaries", level=1)
    for column, summary in context["numeric_summary"].items():
        document.add_paragraph(
            f"{column}: min={summary['min']:.2f}, median={summary['median']:.2f}, "
            f"mean={summary['mean']:.2f}, max={summary['max']:.2f}",
            style="List Bullet",
        )

    document.add_heading("Generated Output Files", level=1)
    for section, values in [
        ("Metadata files", context["metadata_files"]),
        ("Figure files", context["figure_files"]),
        ("Sequence files", context["sequence_files"]),
    ]:
        document.add_paragraph(section, style="List Bullet")
        if values:
            for value in values:
                document.add_paragraph(value, style="List Bullet 2")
        else:
            document.add_paragraph("None", style="List Bullet 2")

    document.save(output_path)
